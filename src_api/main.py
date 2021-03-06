from typing import Optional
import spacy
from fastapi import FastAPI, Form, File, UploadFile
from pydantic import BaseModel
import torch
import torch.nn as nn
import numpy as np
import pandas as pd

import os
import transformers
from transformers import AutoModel, BertTokenizerFast
import shutil
import uvicorn
from docx import Document
from nltk.tokenize import sent_tokenize
import joblib
import  wget
device = torch.device("cpu")

# import BERT-base pretrained model
bert = AutoModel.from_pretrained('bert-base-uncased')

# Load the BERT tokenizer
tokenizer = BertTokenizerFast.from_pretrained('bert-base-uncased')
print("model loaded")
nlp = spacy.load("en_core_web_md")
print("downloading weights from S3")

wget.download("https://test-pytorch-model-demo.s3.ap-south-1.amazonaws.com/weights.pt")
wget.download("https://test-pytorch-model-demo.s3.ap-south-1.amazonaws.com/xb_bst.pkl")
class BERT_Arch(nn.Module):

    def __init__(self, bert):
        super(BERT_Arch, self).__init__()

        self.bert = bert

        # dropout layer
        self.dropout = nn.Dropout(0.1)

        # relu activation function
        self.relu = nn.ReLU()

        # dense layer 1
        self.fc1 = nn.Linear(768, 512)

        # dense layer 2 (Output layer)
        self.fc2 = nn.Linear(512, 6)

        # softmax activation function
        self.softmax = nn.LogSoftmax(dim=1)

    # define the forward pass
    def forward(self, sent_id, mask):
        # pass the inputs to the model
        _, cls_hs = self.bert(sent_id, attention_mask=mask)

        x = self.fc1(cls_hs)

        x = self.relu(x)

        x = self.dropout(x)

        # output layer
        x = self.fc2(x)

        # apply softmax activation
        x = self.softmax(x)

        return x


model = BERT_Arch(bert)

# load weights of best model
path = 'weights.pt'
model.load_state_dict(torch.load(path, map_location=torch.device('cpu')))
print("weights loaded")

app = FastAPI()


@app.post("/agreement_poc")
async def desc_exe(file: UploadFile = File(...)):
    if file.filename.endswith("docx"):
        with open("temp_file.docx", "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    document = Document("./temp_file.docx")
    dummy_text = ""
    agreement_det_para=""
    agreement_on = ""
    agreement_till = ""
    agreement_amount = ""
    result_res ={}
    for p in document.paragraphs:
        dummy_text +=p.text
    sentences = []
    for para in document.paragraphs:
        sentences.append(sent_tokenize(para.text))
    for sent in sentences:
        try:
            sent_id = tokenizer.batch_encode_plus(sent, max_length=500, pad_to_max_length=True, truncation=True)
            check_seq = torch.tensor(sent_id['input_ids'])
            check_mask = torch.tensor(sent_id['attention_mask'])
            check_preds = model(check_seq.to(device), check_mask.to(device))
            check_preds = check_preds.detach().cpu().numpy()
            check_preds_new = np.argmax(check_preds, axis=1)

            sent_text = sent[0]
            print(sent_text,check_preds_new)
            if check_preds_new[0] == 2:
                doc = nlp(sent_text)
                for ent in doc.ents:
                    if ent.label_ == "DATE":
                        agreement_on = agreement_on+ " "+str(ent.text)
            if check_preds_new[0] == 3:
                doc = nlp(sent_text)
                for ent in doc.ents:
                    if ent.label_ == "DATE":
                        agreement_till = agreement_till+ " "+str(ent.text)
            if check_preds_new[0] == 4:
                doc = nlp(sent_text)
                for ent in doc.ents:
                    if ent.label_ == "MONEY":
                        agreement_amount = agreement_amount+" "+str(ent.text)
                    else:
                        result_res["agreement amount "] = " "
            if check_preds_new[0] == 1:
                agreement_det_para = agreement_det_para + " "+ sent_text
        except:
            pass
    result_res["agreement happen on "] = agreement_on
    result_res["agreement happen till "] = agreement_till
    result_res["agreement amount "] = agreement_amount
    result_res["agreement details"] = agreement_det_para
    os.remove("temp_file.docx")
    output_csv = pd.read_csv('OUTPUT.csv')
    output_csv = output_csv.append({"agreement happen on ":agreement_on,"agreement happen till ":agreement_till,
                                    "agreement amount ": agreement_amount,"agreement details": agreement_det_para},
                                   ignore_index=True)
    output_csv.to_csv('OUTPUT.csv', index=False)
    return result_res


@app.post("/fraud_poc")
async def desc_exe(file: UploadFile = File(...)):
    response = {}
    if file.filename.endswith("pkl"):
        with open("temp_df_file.pkl", "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    xg_boost_model = joblib.load("xb_bst.pkl")
    test_data = joblib.load("temp_df_file.pkl")
    result = xg_boost_model.predict(test_data)
    if result[0] == 1:
        response["output"] = "Fraud case"
    elif result[0] == 0:
        response["output"] = "Not Fraud case"
    return response
    os.remove("temp_df_file.pkl")


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
